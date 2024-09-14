import os
import csv
import requests
from io import StringIO
from dotenv import load_dotenv
from telegram import Update, InlineKeyboardButton, InlineKeyboardMarkup
from telegram.ext import Application, CommandHandler, CallbackContext, CallbackQueryHandler, MessageHandler, filters, ConversationHandler
from docx import Document

load_dotenv()

TELEGRAM_API_TOKEN = os.getenv('TELEGRAM_API_TOKEN')
GOOGLE_BOOKS_API_KEY = os.getenv('GOOGLE_BOOKS_API_KEY')

READING_LIST_FILE = 'reading_list.docx'
BOOKS_CSV_FILE = 'books_recommendations.csv'

GENRE, BOOK_NAME, ADD_BOOK, DELETE_BOOK = range(4)

async def start(update: Update, context: CallbackContext) -> None:
    await update.message.reply_text('Welcome to PagePal! Use /help to see available commands.')

async def help_command(update: Update, context: CallbackContext) -> None:
    help_text = (
        "/start - Welcome message\n"
        "/book - Get book recommendations by genre\n"
        "/preview - Get a preview link for a book\n"
        "/list - Manage your reading list\n"
        "/reading_list - View or manage your reading list\n"
        "/help - Show this help message"
    )
    await update.message.reply_text(help_text)

async def book(update: Update, context: CallbackContext) -> int:
    await update.message.reply_text('Please enter the genre you are interested in.')
    return GENRE

async def handle_genre(update: Update, context: CallbackContext) -> None:
    genre = update.message.text
    response = requests.get(f'https://www.googleapis.com/books/v1/volumes?q=subject:{genre}&key={GOOGLE_BOOKS_API_KEY}')
    books = response.json().get('items', [])
    
    if not books:
        await update.message.reply_text('No books found for this genre.')
        return
    
    output = StringIO()
    writer = csv.writer(output)
    writer.writerow(['Title', 'Author', 'Description', 'Year Published', 'Language', 'Preview Link'])

    for book in books:
        volume_info = book.get('volumeInfo', {})
        title = volume_info.get('title', 'N/A')
        authors = ', '.join(volume_info.get('authors', []))
        description = volume_info.get('description', 'N/A')
        published_date = volume_info.get('publishedDate', 'N/A')
        language = volume_info.get('language', 'N/A')
        preview_link = volume_info.get('previewLink', 'N/A')

        writer.writerow([title, authors, description, published_date, language, preview_link])

    output.seek(0)
    await context.bot.send_document(chat_id=update.message.chat_id, document=output.getvalue(), filename=BOOKS_CSV_FILE)
    await update.message.reply_text('Books have been sent as a CSV file.')

async def preview(update: Update, context: CallbackContext) -> int:
    await update.message.reply_text('Please enter the book name.')
    return BOOK_NAME

async def handle_book_name(update: Update, context: CallbackContext) -> None:
    book_name = update.message.text
    response = requests.get(f'https://www.googleapis.com/books/v1/volumes?q=intitle:{book_name}&key={GOOGLE_BOOKS_API_KEY}')
    books = response.json().get('items', [])
    
    if not books:
        await update.message.reply_text('No preview link found for this book.')
        return
    
    volume_info = books[0].get('volumeInfo', {})
    preview_link = volume_info.get('previewLink', 'N/A')

    if preview_link != 'N/A':
        await update.message.reply_text(f'Preview link: {preview_link}')
    else:
        await update.message.reply_text('No preview link available for this book.')

async def list_books(update: Update, context: CallbackContext) -> None:
    keyboard = [
        [InlineKeyboardButton("Add a book", callback_data='add')],
        [InlineKeyboardButton("Delete a book", callback_data='delete')],
        [InlineKeyboardButton("View Reading List", callback_data='view')],
    ]
    reply_markup = InlineKeyboardMarkup(keyboard)
    await update.message.reply_text('Choose an option:', reply_markup=reply_markup)

async def button(update: Update, context: CallbackContext) -> None:
    query = update.callback_query
    data = query.data

    if data == 'add':
        await query.edit_message_text('Send me the book title to add to your reading list.')
        return ADD_BOOK

    if data == 'delete':
        await query.edit_message_text('Send me the book title to remove from your reading list.')
        return DELETE_BOOK

    if data == 'view':
        if os.path.exists(READING_LIST_FILE):
            with open(READING_LIST_FILE, 'rb') as doc:
                await context.bot.send_document(chat_id=update.effective_chat.id, document=doc)
        else:
            await query.edit_message_text('Your reading list is empty.')

async def add_book(update: Update, context: CallbackContext) -> None:
    book_title = update.message.text
    if not os.path.exists(READING_LIST_FILE):
        doc = Document()
    else:
        doc = Document(READING_LIST_FILE)
    
    doc.add_paragraph(f'Title: {book_title}')
    doc.add_paragraph('Preview Link: N/A')
    doc.save(READING_LIST_FILE)
    
    await update.message.reply_text(f'Book "{book_title}" added to your reading list.')

async def delete_book(update: Update, context: CallbackContext) -> None:
    book_title = update.message.text
    if not os.path.exists(READING_LIST_FILE):
        await update.message.reply_text('Reading list is empty.')
        return

    doc = Document(READING_LIST_FILE)
    paragraphs = [p for p in doc.paragraphs if p.text.startswith(f'Title: {book_title}')]
    
    if paragraphs:
        for p in paragraphs:
            p.clear()
        doc.save(READING_LIST_FILE)
        await update.message.reply_text(f'Book "{book_title}" removed from your reading list.')
    else:
        await update.message.reply_text(f'Book "{book_title}" not found in your reading list.')

def main() -> None:
    application = Application.builder().token(TELEGRAM_API_TOKEN).build()

    conv_handler = ConversationHandler(
        entry_points=[CommandHandler('book', book), CommandHandler('preview', preview)],
        states={
            GENRE: [MessageHandler(filters.TEXT & ~filters.COMMAND, handle_genre)],
            BOOK_NAME: [MessageHandler(filters.TEXT & ~filters.COMMAND, handle_book_name)],
            ADD_BOOK: [MessageHandler(filters.TEXT & ~filters.COMMAND, add_book)],
            DELETE_BOOK: [MessageHandler(filters.TEXT & ~filters.COMMAND, delete_book)],
        },
        fallbacks=[]
    )

    application.add_handler(CommandHandler('start', start))
    application.add_handler(CommandHandler('help', help_command))
    application.add_handler(conv_handler)
    application.add_handler(CommandHandler('list', list_books))
    application.add_handler(CommandHandler('reading_list', list_books))
    application.add_handler(CallbackQueryHandler(button))

    application.run_polling()

if __name__ == '__main__':
    main()
